"""
Export generated meeting notes to .docx and .pdf files.
Uses python-docx and reportlab -- both installable via pip, no API key.
"""

import io

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem


def _sections(title, summary, action_items, decisions, keywords, speaker_notes):
    """Normalize inputs into a consistent structure for both exporters."""
    return {
        "title": title or "Meeting Notes",
        "summary": summary or [],
        "action_items": action_items or [],
        "decisions": decisions or [],
        "keywords": keywords or [],
        "speaker_notes": speaker_notes or {},
    }


def export_to_docx(title, summary, action_items, decisions, keywords, speaker_notes) -> bytes:
    data = _sections(title, summary, action_items, decisions, keywords, speaker_notes)

    doc = Document()
    doc.add_heading(data["title"], level=0)

    doc.add_heading("Summary", level=1)
    if data["summary"]:
        for line in data["summary"]:
            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("No summary available.")

    doc.add_heading("Action Items", level=1)
    if data["action_items"]:
        for item in data["action_items"]:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No action items detected.")

    doc.add_heading("Key Decisions", level=1)
    if data["decisions"]:
        for d in data["decisions"]:
            doc.add_paragraph(d, style="List Bullet")
    else:
        doc.add_paragraph("No decisions detected.")

    if data["keywords"]:
        doc.add_heading("Keywords / Topics", level=1)
        doc.add_paragraph(", ".join(data["keywords"]))

    if data["speaker_notes"]:
        doc.add_heading("Speaker-wise Notes", level=1)
        for speaker, lines in data["speaker_notes"].items():
            doc.add_heading(speaker, level=2)
            for line in lines:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(2)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_to_pdf(title, summary, action_items, decisions, keywords, speaker_notes) -> bytes:
    data = _sections(title, summary, action_items, decisions, keywords, speaker_notes)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=LETTER,
                             topMargin=0.75 * inch, bottomMargin=0.75 * inch,
                             leftMargin=0.75 * inch, rightMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    heading_style = styles["Heading1"]
    title_style = ParagraphStyle("TitleStyle", parent=styles["Title"], spaceAfter=14)
    body_style = styles["BodyText"]

    story = [Paragraph(data["title"], title_style), Spacer(1, 8)]

    def add_section(heading, items, empty_msg):
        story.append(Paragraph(heading, heading_style))
        if items:
            story.append(ListFlowable(
                [ListItem(Paragraph(i, body_style)) for i in items],
                bulletType="bullet",
            ))
        else:
            story.append(Paragraph(empty_msg, body_style))
        story.append(Spacer(1, 10))

    add_section("Summary", data["summary"], "No summary available.")
    add_section("Action Items", data["action_items"], "No action items detected.")
    add_section("Key Decisions", data["decisions"], "No decisions detected.")

    if data["keywords"]:
        story.append(Paragraph("Keywords / Topics", heading_style))
        story.append(Paragraph(", ".join(data["keywords"]), body_style))
        story.append(Spacer(1, 10))

    if data["speaker_notes"]:
        story.append(Paragraph("Speaker-wise Notes", heading_style))
        for speaker, lines in data["speaker_notes"].items():
            story.append(Paragraph(speaker, styles["Heading2"]))
            for line in lines:
                story.append(Paragraph(line, body_style))
            story.append(Spacer(1, 6))

    doc.build(story)
    return buffer.getvalue()


def export_to_markdown(title, summary, action_items, decisions, keywords, speaker_notes, **_) -> bytes:
    """Plain Markdown export — handy for pasting into Notion/Obsidian/GitHub.
    Accepts and ignores extra kwargs (e.g. analytics) so callers can pass
    the same results dict used elsewhere without filtering it first."""
    data = _sections(title, summary, action_items, decisions, keywords, speaker_notes)
    lines = [f"# {data['title']}", ""]

    lines.append("## Summary")
    if data["summary"]:
        lines.extend(f"- {line}" for line in data["summary"])
    else:
        lines.append("_No summary available._")
    lines.append("")

    lines.append("## Action Items")
    if data["action_items"]:
        lines.extend(f"- [ ] {item}" for item in data["action_items"])
    else:
        lines.append("_No action items detected._")
    lines.append("")

    lines.append("## Key Decisions")
    if data["decisions"]:
        lines.extend(f"- {d}" for d in data["decisions"])
    else:
        lines.append("_No decisions detected._")
    lines.append("")

    if data["keywords"]:
        lines.append("## Keywords / Topics")
        lines.append(", ".join(f"`{k}`" for k in data["keywords"]))
        lines.append("")

    if data["speaker_notes"]:
        lines.append("## Speaker-wise Notes")
        for speaker, spk_lines in data["speaker_notes"].items():
            lines.append(f"### {speaker}")
            lines.extend(f"- {line}" for line in spk_lines)
        lines.append("")

    return "\n".join(lines).encode("utf-8")
